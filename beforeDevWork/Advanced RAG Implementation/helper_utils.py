# helper_utils.py
import numpy as np
import chromadb
import pandas as pd
from pypdf import PdfReader
import numpy as np
import os
import io
import re
import docx
import html2text


def project_embeddings(embeddings, umap_transform):
    """
    Projects the given embeddings using the provided UMAP transformer.

    Args:
    embeddings (numpy.ndarray): The embeddings to project.
    umap_transform (umap.UMAP): The trained UMAP transformer.

    Returns:
    numpy.ndarray: The projected embeddings.
    """
    projected_embeddings = umap_transform.transform(embeddings)
    return projected_embeddings


def word_wrap(text, width=87):
    """
    Wraps the given text to the specified width.

    Args:
    text (str): The text to wrap.
    width (int): The width to wrap the text to.

    Returns:
    str: The wrapped text.
    """
    return "\n".join([text[i : i + width] for i in range(0, len(text), width)])


def extract_text_from_pdf(file_path):
    """
    Extracts text from a PDF file.

    Args:
    file_path (str): The path to the PDF file.

    Returns:
    str: The extracted text.
    """
    text = []
    with open(file_path, "rb") as f:
        pdf = PdfReader(f)
        for page_num in range(len(pdf.pages)):
            page = pdf.pages[page_num]
            text.append(page.extract_text())
    return "\n".join(text)


def extract_text_from_docx(file_path):
    """
    Extracts text from a DOCX file.

    Args:
    file_path (str): The path to the DOCX file.

    Returns:
    str: The extracted text.
    """
    doc = docx.Document(file_path)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return "\n".join(text)


def extract_text_from_txt(file_path):
    """
    Extracts text from a plain text file with improved encoding detection.
    Supports multiple Chinese encodings including UTF-8, GB18030, Big5, etc.

    Args:
    file_path (str): The path to the text file.

    Returns:
    str: The extracted text.
    """
    # Try different encodings commonly used for Chinese text
    encodings = ['utf-8', 'gb18030', 'big5', 'gbk', 'gb2312', 'utf-16', 'latin1']
    
    # First try to detect if file contains Chinese characters
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding) as f:
                content = f.read(1000)  # Read first 1000 chars to check
                # Check if content contains Chinese characters
                if any('\u4e00' <= c <= '\u9fff' for c in content):
                    # If Chinese characters detected, read the whole file with this encoding
                    with open(file_path, "r", encoding=encoding) as full_f:
                        full_content = full_f.read()
                        print(f"Successfully read file with encoding: {encoding}")
                        return full_content
        except UnicodeDecodeError:
            continue
    
    # If no Chinese encoding works, try to detect encoding automatically
    try:
        import chardet
        with open(file_path, "rb") as f:
            raw_data = f.read()
            result = chardet.detect(raw_data)
            encoding = result['encoding']
            if encoding:
                print(f"Detected encoding: {encoding}")
                return raw_data.decode(encoding)
    except ImportError:
        print("chardet library not found, falling back to default encoding")
    except Exception as e:
        print(f"Error detecting encoding: {str(e)}")
    
    # Fallback to UTF-8 with error replacement
    try:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception as e:
        print(f"Error reading file with fallback encoding: {str(e)}")
        # Last resort - read as binary and decode with replacement
        with open(file_path, "rb") as f:
            return f.read().decode('utf-8', errors='replace')


def extract_text_from_html(file_path):
    """
    Extracts text from an HTML file.

    Args:
    file_path (str): The path to the HTML file.

    Returns:
    str: The extracted text.
    """
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        html_content = f.read()
    
    # Convert HTML to plain text
    h = html2text.HTML2Text()
    h.ignore_links = False
    h.ignore_images = True
    return h.handle(html_content)


def extract_text_from_file(file_path):
    """
    Extracts text from a file based on its extension.

    Args:
    file_path (str): The path to the file.

    Returns:
    str: The extracted text.
    tuple: (text, number of pages/sections)
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == ".pdf":
        text = extract_text_from_pdf(file_path)
        with open(file_path, "rb") as f:
            pdf = PdfReader(f)
            num_pages = len(pdf.pages)
        return text, num_pages
    
    elif file_extension == ".docx":
        text = extract_text_from_docx(file_path)
        doc = docx.Document(file_path)
        num_paragraphs = len(doc.paragraphs)
        return text, num_paragraphs
    
    elif file_extension == ".txt":
        text = extract_text_from_txt(file_path)
        # Count paragraphs in text file (separated by blank lines)
        paragraphs = re.split(r'\n\s*\n', text)
        return text, len(paragraphs)
    
    elif file_extension in [".html", ".htm"]:
        text = extract_text_from_html(file_path)
        # Count paragraphs in HTML (rough estimate)
        paragraphs = re.split(r'\n\s*\n', text)
        return text, len(paragraphs)
    
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")


def load_chroma(filename, collection_name, embedding_function):
    """
    Loads a document from a file, extracts text, generates embeddings, and stores it in a Chroma collection.

    Args:
    filename (str): The path to the file.
    collection_name (str): The name of the Chroma collection.
    embedding_function (callable): A function to generate embeddings.

    Returns:
    chroma.Collection: The Chroma collection with the document embeddings.
    """
    # Extract text from the file
    text, _ = extract_text_from_file(filename)

    # Split text into paragraphs or chunks
    paragraphs = text.split("\n\n")

    # Generate embeddings for each chunk
    embeddings = [embedding_function(paragraph) for paragraph in paragraphs]

    # Create a DataFrame to store text and embeddings
    data = {"text": paragraphs, "embeddings": embeddings}
    df = pd.DataFrame(data)

    # Create or load the Chroma collection
    collection = chromadb.Client().create_collection(collection_name)

    # Add the data to the Chroma collection
    for ids, row in df.iterrows():
        collection.add(ids=ids, documents=row["text"], embeddings=row["embeddings"])

    return collection
